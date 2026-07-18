from __future__ import annotations

import math
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[2]
MANUAL_ROOT = (
    REPO_ROOT
    / "gann-astro-desk"
    / "docs"
    / "user_manual"
    / "sarvatobhadra_chakra"
)
ASSET_ROOT = MANUAL_ROOT / "assets"
QA_ROOT = MANUAL_ROOT / "qa_render"
SOURCE_SCREENSHOT = ASSET_ROOT / "chakra_lab_native_0109.jpg"
OUTPUT_DOCX = MANUAL_ROOT / "Gann_Astro_Desk_Sarvatobhadra_Chakra_Manual_v1.0.docx"
OUTPUT_PDF = MANUAL_ROOT / "Gann_Astro_Desk_Sarvatobhadra_Chakra_Manual_v1.0.pdf"
QUICK_START_PNG = MANUAL_ROOT / "Sarvatobhadra_Chakra_Quick_Start_v1.0.png"
ANNOTATED_SCREEN_PNG = ASSET_ROOT / "chakra_lab_native_0109_annotated.png"
RAY_GEOMETRY_PNG = ASSET_ROOT / "vedha_ray_geometry.png"

FONT_REGULAR = Path(r"C:\Windows\Fonts\segoeui.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\segoeuib.ttf")
FONT_MONO = Path(r"C:\Windows\Fonts\consola.ttf")

NAVY = "173A4A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "2E8E7B"
TEAL_LIGHT = "E8F5F1"
GOLD = "B78125"
GOLD_LIGHT = "FFF4D6"
MAGENTA = "A34F8B"
MAGENTA_LIGHT = "F7EAF3"
CYAN = "3899BB"
CYAN_LIGHT = "E7F5FA"
INK = "17242D"
MUTED = "5D6A73"
LIGHT = "F2F5F7"
BORDER = "CCD5DB"
RED = "A43B3B"
RED_LIGHT = "FCECEC"
WHITE = "FFFFFF"


def font(size: int, bold: bool = False, mono: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_MONO if mono else FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(str(path), size)


def rgb(hex_color: str) -> tuple[int, int, int]:
    value = hex_color.lstrip("#")
    return tuple(int(value[index : index + 2], 16) for index in (0, 2, 4))


def wrap_text(draw: ImageDraw.ImageDraw, text: str, face: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=face)[2] <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    face: ImageFont.FreeTypeFont,
    fill: tuple[int, int, int],
    width: int,
    spacing: int = 6,
) -> int:
    x, y = xy
    lines = wrap_text(draw, text, face, width)
    line_height = face.size + spacing
    for line in lines:
        draw.text((x, y), line, font=face, fill=fill)
        y += line_height
    return y


def create_quick_start_infographic() -> None:
    canvas = Image.new("RGB", (1600, 900), rgb("F4F7F8"))
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, 1600, 118), fill=rgb("11202A"))
    draw.text((64, 30), "SARVATOBHADRA CHAKRA LAB", font=font(42, bold=True), fill=rgb(WHITE))
    draw.text(
        (66, 82),
        "Six-step beginner workflow in Gann Astro Desk 0.10.9",
        font=font(20),
        fill=rgb("B8C6CE"),
    )

    steps = [
        ("1", "Open Chakra", "Choose Chakra in the top workspace bar."),
        ("2", "Set moment and place", "Enter the IST timestamp and intentional research coordinates."),
        ("3", "Choose actors", "For a first run, keep fixed actors and uncheck planets marked Required."),
        ("4", "Refresh snapshot", "Inputs do not recalculate until you press Refresh snapshot."),
        ("5", "Read the board", "Green = context, gold = ray, magenta = matched, cyan = selected."),
        ("6", "Read the ledger", "Check score, favorable/adverse units, coverage, matched cells, and inspector."),
    ]
    card_w = 468
    card_h = 260
    gap_x = 34
    gap_y = 30
    start_x = 64
    start_y = 155
    accents = [TEAL, BLUE, GOLD, MAGENTA, CYAN, DARK_BLUE]

    for index, (number, title, detail) in enumerate(steps):
        row, column = divmod(index, 3)
        x = start_x + column * (card_w + gap_x)
        y = start_y + row * (card_h + gap_y)
        accent = rgb(accents[index])
        draw.rounded_rectangle(
            (x, y, x + card_w, y + card_h),
            radius=14,
            fill=rgb(WHITE),
            outline=rgb(BORDER),
            width=2,
        )
        draw.rounded_rectangle(
            (x, y, x + 84, y + card_h),
            radius=14,
            fill=accent,
        )
        draw.rectangle((x + 68, y, x + 84, y + card_h), fill=accent)
        draw.ellipse((x + 17, y + 24, x + 67, y + 74), fill=rgb(WHITE))
        number_box = draw.textbbox((0, 0), number, font=font(27, bold=True))
        number_w = number_box[2] - number_box[0]
        draw.text(
            (x + 42 - number_w / 2, y + 32),
            number,
            font=font(27, bold=True),
            fill=accent,
        )
        draw.text((x + 112, y + 31), title, font=font(26, bold=True), fill=rgb(INK))
        draw_wrapped(
            draw,
            (x + 112, y + 89),
            detail,
            font(22),
            rgb(MUTED),
            card_w - 142,
            spacing=9,
        )

    draw.rounded_rectangle(
        (64, 755, 1536, 850),
        radius=12,
        fill=rgb(RED_LIGHT),
        outline=rgb("DCA5A5"),
        width=2,
    )
    draw.text((90, 777), "STOP", font=font(23, bold=True), fill=rgb(RED))
    draw_wrapped(
        draw,
        (180, 773),
        "The percentage is an evidence-balance meter. It is not probability, confidence, "
        "bullish/bearish direction, or permission to trade.",
        font(22, bold=True),
        rgb("6B2727"),
        1310,
        spacing=7,
    )
    canvas.save(QUICK_START_PNG, quality=95)


def create_annotated_screenshot() -> None:
    if not SOURCE_SCREENSHOT.is_file():
        raise FileNotFoundError(SOURCE_SCREENSHOT)
    source = Image.open(SOURCE_SCREENSHOT).convert("RGB")
    canvas = Image.new("RGB", (source.width, source.height + 150), rgb("0F171F"))
    canvas.paste(source, (0, 0))
    draw = ImageDraw.Draw(canvas, "RGBA")

    targets = [
        (1, (410, 47), TEAL),
        (2, (140, 185), BLUE),
        (3, (140, 345), GOLD),
        (4, (140, 555), MAGENTA),
        (5, (760, 430), CYAN),
        (6, (1125, 124), GOLD),
        (7, (1365, 165), TEAL),
        (8, (1370, 550), MAGENTA),
    ]
    for number, (x, y), color in targets:
        accent = (*rgb(color), 255)
        draw.ellipse((x - 22, y - 22, x + 22, y + 22), fill=(15, 23, 31, 220), outline=accent, width=5)
        label = str(number)
        box = draw.textbbox((0, 0), label, font=font(24, bold=True))
        draw.text(
            (x - (box[2] - box[0]) / 2, y - 16),
            label,
            font=font(24, bold=True),
            fill=(255, 255, 255, 255),
        )

    legend = [
        ("1", "Chakra tab"),
        ("2", "Moment and place"),
        ("3", "Optional context"),
        ("4", "Vedha actors"),
        ("5", "81-cell board"),
        ("6", "Board legend"),
        ("7", "Guidance ledger"),
        ("8", "Evidence and inspector"),
    ]
    legend_y = source.height + 20
    column_width = source.width // 4
    for index, (number, label) in enumerate(legend):
        row, column = divmod(index, 4)
        x = column * column_width + 28
        y = legend_y + row * 58
        draw.ellipse((x, y, x + 34, y + 34), fill=(*rgb(TEAL), 255))
        draw.text((x + 11, y + 4), number, font=font(18, bold=True), fill=(255, 255, 255, 255))
        draw.text((x + 48, y + 5), label, font=font(20, bold=True), fill=(225, 233, 238, 255))
    canvas.save(ANNOTATED_SCREEN_PNG)


def draw_ray_grid(
    draw: ImageDraw.ImageDraw,
    origin: tuple[int, int],
    label: str,
    direction: str,
    accent: str,
) -> None:
    ox, oy = origin
    cell = 42
    size = cell * 9
    draw.text((ox, oy - 55), label, font=font(27, bold=True), fill=rgb(INK))
    for row in range(9):
        for column in range(9):
            x1 = ox + column * cell
            y1 = oy + row * cell
            draw.rectangle((x1, y1, x1 + cell, y1 + cell), fill=rgb("FAFBFC"), outline=rgb("CAD3D9"), width=1)
    source = (0, 4)
    source_center = (ox + (source[1] + 0.5) * cell, oy + 0.5 * cell)
    draw.rectangle(
        (
            ox + source[1] * cell + 2,
            oy + 2,
            ox + (source[1] + 1) * cell - 2,
            oy + cell - 2,
        ),
        fill=rgb(accent),
    )
    draw.text(
        (source_center[0] - 9, source_center[1] - 12),
        "S",
        font=font(19, bold=True),
        fill=rgb(WHITE),
    )
    if direction == "LEFT":
        points = [(row, 4 - row) for row in range(1, 5)]
    elif direction == "RIGHT":
        points = [(row, 4 + row) for row in range(1, 5)]
    else:
        points = [(8, 4)]
    line_points = [source_center]
    for row, column in points:
        cx = ox + (column + 0.5) * cell
        cy = oy + (row + 0.5) * cell
        line_points.append((cx, cy))
        draw.rectangle(
            (
                ox + column * cell + 4,
                oy + row * cell + 4,
                ox + (column + 1) * cell - 4,
                oy + (row + 1) * cell - 4,
            ),
            fill=rgb(GOLD_LIGHT),
            outline=rgb(GOLD),
            width=3,
        )
    draw.line(line_points, fill=rgb(accent), width=6)
    end_x, end_y = line_points[-1]
    angle = math.atan2(end_y - line_points[-2][1], end_x - line_points[-2][0])
    arrow = []
    for offset in (2.6, -2.6):
        arrow.append(
            (
                end_x - 18 * math.cos(angle + offset),
                end_y - 18 * math.sin(angle + offset),
            )
        )
    draw.polygon([(end_x, end_y), arrow[0], arrow[1]], fill=rgb(accent))


def create_ray_geometry_infographic() -> None:
    canvas = Image.new("RGB", (1600, 650), rgb("F4F7F8"))
    draw = ImageDraw.Draw(canvas)
    draw.text((60, 28), "HOW THE PROGRAM DRAWS VEDHA RAYS", font=font(35, bold=True), fill=rgb(INK))
    draw.text(
        (62, 78),
        "The direction is relative to a planet facing inward from its nakshatra cell.",
        font=font(21),
        fill=rgb(MUTED),
    )
    draw_ray_grid(draw, (60, 170), "LEFT", "LEFT", TEAL)
    draw_ray_grid(draw, (610, 170), "FRONT", "FRONT", BLUE)
    draw_ray_grid(draw, (1160, 170), "RIGHT", "RIGHT", MAGENTA)
    draw.rounded_rectangle((60, 575, 1540, 630), radius=10, fill=rgb(GOLD_LIGHT), outline=rgb("D7B96D"), width=2)
    draw.text(
        (82, 590),
        "Figure-relative only: LEFT and RIGHT are not compass directions. Absolute cardinal orientation remains unresolved.",
        font=font(20, bold=True),
        fill=rgb("6B531C"),
    )
    canvas.save(RAY_GEOMETRY_PNG)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches: Iterable[float], indent_dxa: int = 120) -> None:
    widths = list(widths_inches)
    total_dxa = sum(int(round(value * 1440)) for value in widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(grid_col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths)):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(round(width * 1440))))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size: float, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Subtitle", 15, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    for document_section in doc.sections:
        header = document_section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run("GANN ASTRO DESK  |  SARVATOBHADRA CHAKRA LAB")
        set_run_font(run, 8.5, MUTED, bold=True)
        footer = document_section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_paragraph.paragraph_format.space_before = Pt(2)
        footer_run = footer_paragraph.add_run("USER MANUAL  |  PAGE ")
        set_run_font(footer_run, 8.5, MUTED)
        add_field(footer_paragraph, "PAGE")


def add_kicker(doc: Document, text: str, color: str = TEAL) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(text.upper())
    set_run_font(run, 10, color, bold=True)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, 11, INK, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest, 11, INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 11, INK)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, 11, INK)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run, 11, INK)


def add_callout(
    doc: Document,
    label: str,
    text: str,
    *,
    fill: str = TEAL_LIGHT,
    border: str = TEAL,
    label_color: str = TEAL,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_node = OxmlElement(f"w:{edge}")
        edge_node.set(qn("w:val"), "single")
        edge_node.set(qn("w:sz"), "8")
        edge_node.set(qn("w:color"), border)
        borders.append(edge_node)
    set_repeat_table_header(table.rows[0])
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, 10.5, label_color, bold=True)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, 10.5, INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_picture(doc: Document, path: Path, width: float, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    inline_shape._inline.docPr.set("descr", caption)
    inline_shape._inline.docPr.set("title", path.stem.replace("_", " "))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(3)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, 9, MUTED, italic=True)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    header_fill: str = "E8EEF5",
    accent_first_column: bool = False,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_shading(header_cells[index], header_fill)
        paragraph = header_cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, 9.5, INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            if accent_first_column and index == 0:
                set_cell_shading(cells[index], LIGHT)
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, 9.5, INK, bold=accent_first_column and index == 0)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    run = paragraph.add_run()
    set_run_font(run, 1, WHITE)


def build_document() -> None:
    doc = Document()
    configure_document(doc)

    add_kicker(doc, "Read-only research guide", GOLD)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Sarvatobhadra Chakra Lab")
    set_run_font(title_run, 30, NAVY, bold=True)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Easy user manual for Gann Astro Desk 0.10.9")
    set_run_font(subtitle_run, 15, MUTED)

    doc.add_paragraph()
    add_picture(
        doc,
        RAY_GEOMETRY_PNG,
        6.25,
        "The lab uses source-profiled, figure-relative Vedha geometry on an 81-cell board.",
    )
    badge_table = doc.add_table(rows=1, cols=3)
    badge_data = [
        ("READ ONLY", TEAL_LIGHT, TEAL),
        ("NO LOOKAHEAD", CYAN_LIGHT, CYAN),
        ("NOT FINANCIALLY VALIDATED", GOLD_LIGHT, GOLD),
    ]
    for index, (text, fill, color) in enumerate(badge_data):
        cell = badge_table.cell(0, index)
        set_cell_shading(cell, fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, 9, color, bold=True)
    set_repeat_table_header(badge_table.rows[0])
    set_table_geometry(badge_table, [2.0, 2.0, 2.5], indent_dxa=0)
    doc.add_paragraph()
    add_callout(
        doc,
        "Purpose",
        "Use this lab to inspect source-profiled Sarvatobhadra context and Vedha evidence for one moment. "
        "Its percentage is an evidence-balance meter, not probability, confidence, market direction, or permission to trade.",
        fill=RED_LIGHT,
        border=RED,
        label_color=RED,
    )
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_run = metadata.add_run("Manual version 1.0  |  18 July 2026  |  Raman sidereal profile")
    set_run_font(metadata_run, 9.5, MUTED)

    page_break(doc)
    doc.add_heading("Quick Start", level=1)
    add_picture(
        doc,
        QUICK_START_PNG,
        6.5,
        "Begin with this six-step workflow. Detailed explanations follow.",
    )
    add_callout(
        doc,
        "Best first run",
        "Keep Sun, Moon, Rahu, and Ketu selected. Uncheck Mars, Mercury, Jupiter, Venus, and Saturn "
        "until you have verified their motion class. This avoids guessing.",
        fill=GOLD_LIGHT,
        border=GOLD,
        label_color=GOLD,
    )

    page_break(doc)
    doc.add_heading("Screen Tour", level=1)
    add_picture(
        doc,
        ANNOTATED_SCREEN_PNG,
        6.5,
        "Actual native 0.10.9 Chakra Lab screen, annotated for this manual.",
    )
    tour_rows = [
        ["1", "Chakra tab", "Opens the read-only Chakra Lab workspace."],
        ["2", "Moment and place", "IST timestamp, latitude, longitude, and altitude."],
        ["3", "Optional context", "Research-only vowel and name-initial keys."],
        ["4", "Vedha actors", "Planet selection, explicit motion, and dignity."],
        ["5", "81-cell board", "Nakshatra, rashi, tithi, weekday, vowel, and initial layers."],
        ["6", "Board legend", "Context, ray, and matched color states."],
        ["7", "Guidance ledger", "Evidence balance, units, net, and coverage."],
        ["8", "Evidence and inspector", "Actor readiness, matched cells, and selected-cell layers."],
    ]
    add_table(doc, ["No.", "Area", "What it is for"], tour_rows, [0.55, 1.55, 4.4], accent_first_column=True)

    page_break(doc)
    doc.add_heading("The Idea in One Minute", level=1)
    add_body(
        doc,
        "The program builds one 9 x 9 Sarvatobhadra board for the moment you choose. "
        "Each cell can contain one or more certified layers, such as a nakshatra, rashi, "
        "tithi group, weekday, vowel, or name-initial key."
    )
    add_numbered(doc, "The current moment supplies context values: planetary nakshatras and rashis, the tithi group, and any optional letter keys you enter.")
    add_numbered(doc, "Each selected planet starts from its current nakshatra cell and projects one source-profiled Vedha ray.")
    add_numbered(doc, "A matched cell occurs only where a ray target overlaps the current context.")
    add_numbered(doc, "Only matched cells enter the guidance ledger; ordinary ray cells do not score.")
    add_picture(
        doc,
        RAY_GEOMETRY_PNG,
        6.5,
        "Left, front, and right are relative to the figure. They are not North, South, East, or West.",
    )
    add_callout(
        doc,
        "Important",
        "The program deliberately does not claim an absolute compass orientation because the held source figures have an unresolved rotation conflict.",
        fill=GOLD_LIGHT,
        border=GOLD,
        label_color=GOLD,
    )

    page_break(doc)
    doc.add_heading("1. Set the Moment", level=1)
    doc.add_heading("Timestamp", level=2)
    add_body(
        doc,
        "Enter the moment in IST. The app sends an explicit +05:30 offset and derives astronomy, "
        "Panchanga, placements, board context, and Vedha evidence from the same cutoff. "
        "There is no lookahead inside the snapshot."
    )
    add_bullet(doc, "For current research, use the present IST date and time.")
    add_bullet(doc, "For historical research, enter the exact historical IST moment you intend to study.")
    add_bullet(doc, "After changing any input, press Refresh snapshot. Inputs do not recalculate automatically.")
    doc.add_heading("Location", level=2)
    add_body(
        doc,
        "Latitude, longitude, and altitude define the geographic reference used by the foundation calculation. "
        "The fields are prefilled from the active workspace; verify them instead of assuming they represent your laptop location."
    )
    add_callout(
        doc,
        "Consistency rule",
        "Use one documented location policy for a research series. Changing location between samples can change Panchanga context and makes comparisons harder to interpret.",
    )
    doc.add_heading("Optional context", level=2)
    add_body(
        doc,
        "Vowel keys and name-initial keys add extra context layers. Leave both blank for a basic moment study. "
        "Do not type an instrument symbol such as USDJPY unless you have a separately reviewed, time-valid mapping."
    )
    add_table(
        doc,
        ["Field", "Example", "When to use"],
        [
            ["Vowel keys", "A, AA", "Only when a source-reviewed vowel mapping is part of the research question."],
            ["Name-initial keys", "KA, RA", "Only when a source-reviewed name or instrument identity mapping exists."],
        ],
        [1.4, 1.2, 3.9],
        accent_first_column=True,
    )
    add_callout(
        doc,
        "Current boundary",
        "The separate instrument-relative Forex research layer is not connected to the production Chakra Lab. Optional letters are therefore manual research inputs, not automatic currency mappings.",
        fill=GOLD_LIGHT,
        border=GOLD,
        label_color=GOLD,
    )

    page_break(doc)
    doc.add_heading("2. Choose Vedha Actors", level=1)
    add_body(
        doc,
        "A Vedha actor is a selected planet whose current nakshatra becomes the starting cell for one ray. "
        "The checkbox includes or excludes the actor from the guidance calculation."
    )
    doc.add_heading("Direction rules used by the program", level=2)
    direction_rows = [
        ["Sun", "Fixed", "LEFT", "No motion choice is needed."],
        ["Moon", "Fixed", "LEFT", "Waxing/waning context resolves nature automatically."],
        ["Rahu", "Fixed", "RIGHT", "Treated as fixed source-profile direction."],
        ["Ketu", "Fixed", "RIGHT", "Treated as fixed source-profile direction."],
        ["Mars, Mercury, Jupiter, Venus, Saturn", "Direct / swift", "LEFT", "Choose only with verified motion evidence."],
        ["Mars, Mercury, Jupiter, Venus, Saturn", "Mean", "FRONT", "Choose only with verified motion evidence."],
        ["Mars, Mercury, Jupiter, Venus, Saturn", "Retrograde", "RIGHT", "Choose only with verified motion evidence."],
    ]
    add_table(doc, ["Actor", "Motion", "Ray", "Beginner instruction"], direction_rows, [1.7, 1.25, 0.75, 2.8])
    add_callout(
        doc,
        "Why the app says Required",
        "Automatic direct-versus-swift thresholds are not certified. The program refuses to guess. "
        "If you do not have verified motion evidence, uncheck that planet or leave it unresolved.",
        fill=GOLD_LIGHT,
        border=GOLD,
        label_color=GOLD,
    )
    doc.add_heading("Nature rules", level=2)
    nature_rows = [
        ["Jupiter, Venus", "Benefic", "Matched evidence receives a positive sign."],
        ["Saturn, Sun, Rahu, Ketu, Mars", "Malefic", "Matched evidence receives a negative sign."],
        ["Moon", "Conditional", "Non-waning is benefic; waning is malefic in this profile."],
        ["Mercury", "Conditional", "Needs association context; the current UI does not infer it, so matched Mercury evidence can remain unresolved."],
    ]
    add_table(doc, ["Bodies", "Nature", "Program behavior"], nature_rows, [2.2, 1.15, 3.15], accent_first_column=True)

    doc.add_heading("3. Set Dignity Without Double Counting", level=1)
    add_body(
        doc,
        "Dignity changes the size of a scored match. It does not change the ray direction. "
        "Use the status that is valid for the selected moment; do not choose a value merely to make the score larger."
    )
    add_table(
        doc,
        ["Setting", "Multiplier", "Meaning in this ledger"],
        [
            ["Ordinary", "1.0x", "One ordinary matched evidence unit."],
            ["Retrograde", "2.0x", "Applied from the actor's motion class."],
            ["Exalted", "3.0x", "Three evidence units for an otherwise resolved match."],
            ["Debilitated", "0.5x", "Half an evidence unit for an otherwise resolved match."],
        ],
        [1.5, 1.1, 3.9],
        accent_first_column=True,
    )
    add_callout(
        doc,
        "Fail-closed combination",
        "Do not combine Retrograde with Exalted or Debilitated. The held source gives separate multipliers but no certified stacking or precedence rule. "
        "The program marks that contribution unresolved and excludes it from the net.",
        fill=RED_LIGHT,
        border=RED,
        label_color=RED,
    )
    add_body(
        doc,
        "A contribution can also be unresolved when the planet's nature is conditional, such as Mercury without certified association context. "
        "Unresolved matches remain visible and reduce Coverage; they are not silently treated as zero evidence."
    )

    page_break(doc)
    doc.add_heading("4. Read the 81-Cell Board", level=1)
    add_table(
        doc,
        ["Visual state", "Meaning", "What to do"],
        [
            ["Green - Context", "The cell contains a value active in the selected moment.", "Treat it as current context, not a signal."],
            ["Gold outline - Ray", "At least one selected actor's Vedha ray reaches the cell.", "Inspect only if it overlaps context."],
            ["Magenta - Matched", "A ray target and current context overlap.", "This match appears in Matched cells and may contribute units."],
            ["Cyan outline - Selected", "You clicked the cell.", "Read every certified layer in Cell inspector."],
        ],
        [1.45, 2.35, 2.7],
        accent_first_column=True,
    )
    add_body(
        doc,
        "A cell can hold several layers at once. The large label shows the preferred display layer, while the small abbreviations indicate all layers present. "
        "Clicking a cell is the reliable way to see its complete contents."
    )
    add_numbered(doc, "Look for magenta matched cells first.")
    add_numbered(doc, "Click a matched cell or its entry in Matched cells.")
    add_numbered(doc, "Read the layer and value in Cell inspector.")
    add_numbered(doc, "Check which actor produced the match and whether the contribution was scored or unresolved.")
    add_callout(
        doc,
        "No matched cells is valid",
        "It means the selected rays did not intersect the current target context. It is not a calculation failure and it is not a neutral market forecast.",
    )

    page_break(doc)
    doc.add_heading("5. Read the Guidance Ledger", level=1)
    add_body(
        doc,
        "The ledger summarizes only matched target layers. Its numerical model is an explicit engineering normalization built for comparison. "
        "It is not a classical numerical score."
    )
    add_table(
        doc,
        ["Metric", "Plain-English meaning"],
        [
            ["Favorable", "Sum of positive, resolved matched evidence units."],
            ["Adverse", "Sum of negative, resolved matched evidence units."],
            ["Net", "Favorable plus Adverse."],
            ["Coverage", "Scored matched cells divided by all matched cells."],
            ["Percentage", "Net divided by Favorable plus absolute Adverse, shown from -100% to +100%."],
            ["Band", "Favorable dominant, adverse dominant, or balanced/no scored hits."],
        ],
        [1.45, 5.05],
        accent_first_column=True,
    )
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(10)
    formula.paragraph_format.space_after = Pt(10)
    formula_run = formula.add_run("normalized score = net / (favorable + absolute adverse)")
    formula_run.font.name = "Consolas"
    formula_run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
    formula_run.font.size = Pt(11)
    formula_run.font.color.rgb = RGBColor.from_string(NAVY)
    formula_run.bold = True
    add_callout(
        doc,
        "Do not misread -100%",
        "It can mean one adverse match and no favorable match. It does not mean 100% confidence, a 100% chance of decline, or a maximum-strength event. "
        "Always read matched count and coverage beside the percentage.",
        fill=RED_LIGHT,
        border=RED,
        label_color=RED,
    )
    add_callout(
        doc,
        "Coverage rule",
        "A low coverage percentage means some matched evidence could not be scored because motion, nature, or multiplier precedence was unresolved. "
        "Do not compare low-coverage samples as if they were complete.",
        fill=GOLD_LIGHT,
        border=GOLD,
        label_color=GOLD,
    )

    page_break(doc)
    doc.add_heading("Worked Example", level=1)
    add_body(
        doc,
        "This engine-level example is intentionally simple. It demonstrates the arithmetic, not a market forecast."
    )
    add_table(
        doc,
        ["Actor", "Source and ray", "Matched target", "Signed units"],
        [
            ["Jupiter", "Krittika, Mean -> Front", "Shravana", "+1.0"],
            ["Saturn", "Krittika, Retrograde -> Right", "Bharani", "-2.0"],
        ],
        [1.25, 2.4, 1.45, 1.4],
        accent_first_column=True,
    )
    add_body(doc, "Favorable = +1.0. Adverse = -2.0. Net = -1.0. Absolute evidence = 3.0.")
    result = doc.add_paragraph()
    result.alignment = WD_ALIGN_PARAGRAPH.CENTER
    result_run = result.add_run("-1.0 / 3.0 = -33.3%  |  Adverse Evidence Dominant")
    set_run_font(result_run, 15, RED, bold=True)
    add_callout(
        doc,
        "Correct interpretation",
        "Within the supplied Sarvatobhadra context, resolved adverse evidence outweighed resolved favorable evidence. "
        "This does not mean short USDJPY, sell Bitcoin, or place any order.",
    )
    add_callout(
        doc,
        "Incorrect interpretation",
        "There is a 33.3% probability of a price decline, or Saturn guarantees a bearish move.",
        fill=RED_LIGHT,
        border=RED,
        label_color=RED,
    )

    page_break(doc)
    doc.add_heading("Recommended Beginner Workflows", level=1)
    doc.add_heading("Workflow A - Safe observation run", level=2)
    add_numbered(doc, "Open Chakra and verify the IST timestamp and location.")
    add_numbered(doc, "Keep Sun, Moon, Rahu, and Ketu selected.")
    add_numbered(doc, "Uncheck Mars, Mercury, Jupiter, Venus, and Saturn.")
    add_numbered(doc, "Leave vowel and name-initial fields blank.")
    add_numbered(doc, "Press Refresh snapshot.")
    add_numbered(doc, "Record snapshot ID, score, coverage, matched cells, and the fact that only fixed actors were used.")
    doc.add_heading("Workflow B - Full research run", level=2)
    add_numbered(doc, "Prepare verified motion classes for Mars through Saturn before opening the lab.")
    add_numbered(doc, "Select only the actors relevant to the research question.")
    add_numbered(doc, "Set dignity from the same timestamp, avoiding retrograde plus non-ordinary combinations.")
    add_numbered(doc, "Refresh and inspect Actor evidence for Ready or unresolved status.")
    add_numbered(doc, "Record every actor setting together with the snapshot ID.")
    doc.add_heading("Workflow C - Instrument identity experiment", level=2)
    add_body(
        doc,
        "Use optional vowel/name-initial keys only with a reviewed mapping record. Keep this work in the separate experimental instrument-relative SBC layer. "
        "Do not treat manual letter matching as production financial evidence."
    )

    page_break(doc)
    doc.add_heading("Troubleshooting", level=1)
    troubleshooting_rows = [
        ["Motion Required", "A selected variable planet has no explicit motion class.", "Set a verified class or uncheck the actor; do not guess."],
        ["No matched target cells", "No ray intersects current context.", "Record zero matches; do not call it neutral or bullish/bearish."],
        ["0.0% score", "Evidence is balanced or there are no scored hits.", "Check matched count, favorable/adverse units, and coverage."],
        ["Low coverage", "One or more matches are unresolved.", "Inspect Mercury nature, missing motion, and multiplier conflicts."],
        ["Score did not change", "Inputs were edited but the snapshot was not refreshed.", "Press Refresh snapshot and confirm the snapshot ID changes."],
        ["Unexpected time", "Timestamp was entered incorrectly.", "Use IST and verify date, hour, and minute before refreshing."],
        ["Invalid vowel or initial", "The token is outside the certified board vocabulary.", "Use the accepted token lists in this manual."],
        ["Not financially validated", "This is the correct permanent research lock.", "Do not seek a setting that removes it."],
    ]
    add_table(doc, ["What you see", "What it means", "What to do"], troubleshooting_rows, [1.55, 2.25, 2.7])
    add_callout(
        doc,
        "Snapshot identity",
        "The short code beside Refresh snapshot identifies the exact inputs and evidence. If inputs change, the ID should change. Save it with your research note.",
    )

    page_break(doc)
    doc.add_heading("Research Note Template", level=1)
    add_body(
        doc,
        "Use this checklist whenever you compare Chakra evidence with a later market observation. "
        "Keeping the input record separate from the future outcome protects against hindsight leakage."
    )
    note_rows = [
        ["Snapshot ID", "Copy the short ID shown at the top."],
        ["Moment", "IST date and time used."],
        ["Location", "Latitude, longitude, altitude, and reason for choosing them."],
        ["Context keys", "Vowels and initials, or 'blank'."],
        ["Actors", "Selected planets plus exact motion and dignity settings."],
        ["Ledger", "Percentage, band, favorable, adverse, net, coverage."],
        ["Matched cells", "Actor, layer, value, direction, and signed units/unresolved status."],
        ["Pre-outcome hypothesis", "A plain statement written before seeing later bars."],
        ["Outcome observation", "Recorded later in a separate timestamped field."],
        ["Decision", "Keep as research evidence; never backfill it into live inference."],
    ]
    add_table(doc, ["Field", "What to record"], note_rows, [1.75, 4.75], accent_first_column=True)
    add_callout(
        doc,
        "Recommended language",
        "Use 'favorable/adverse evidence dominant' exactly as the app does. Avoid replacing it with bullish, bearish, buy, sell, confidence, or probability.",
        fill=GOLD_LIGHT,
        border=GOLD,
        label_color=GOLD,
    )

    page_break(doc)
    doc.add_heading("Accepted Optional Keys", level=1)
    add_body(
        doc,
        "The program uppercases comma-separated input and rejects unknown values. The following machine tokens are accepted by the current 81-cell profile."
    )
    add_table(
        doc,
        ["Layer", "Accepted keys"],
        [
            [
                "Vowels",
                "A, AA, I, II, U, UU, VOCALIC_R, LONG_VOCALIC_R, VOCALIC_L, "
                "LONG_VOCALIC_L, E, AI, O, AU, ANUSVARA, VISARGA",
            ],
            [
                "Name initials",
                "A, VA, KA, HA, DDA, MA, TTA, PA, RA, TA, NA, YA, BHA, JA, "
                "KHA, GA, SA, DA, CHA, LA",
            ],
        ],
        [1.35, 5.15],
        accent_first_column=True,
    )
    add_callout(
        doc,
        "Semantic exception",
        "NAME_INITIAL:A is kept as its own board cell because the source ring begins with vowel A. The engine does not create a duplicate synthetic vowel target.",
    )
    doc.add_heading("Plain-English glossary", level=2)
    glossary_rows = [
        ["Actor", "A selected planet that projects one Vedha ray."],
        ["Context", "Values active in the selected moment."],
        ["Ray", "Figure-relative path generated from an actor's nakshatra cell."],
        ["Matched cell", "A ray target that overlaps current context."],
        ["Contribution", "One matched layer with resolved or unresolved guidance units."],
        ["Coverage", "Share of matched contributions that could be scored."],
        ["Evidence cutoff", "The same timestamp used for all facts in one snapshot."],
        ["Guidance band", "Which sign dominates the resolved matched evidence."],
    ]
    add_table(doc, ["Term", "Meaning"], glossary_rows, [1.5, 5.0], accent_first_column=True)

    page_break(doc)
    doc.add_heading("What the Chakra Lab Does Not Do", level=1)
    blocked = [
        "It does not read price bars or MT5 market data.",
        "It does not produce bullish, bearish, buy, sell, stop, target, position size, or P/L.",
        "It does not feed Auto Suggest, the prospective shadow ledger, or order execution.",
        "It does not infer direct/swift versus mean motion for Mars through Saturn.",
        "It does not infer Mercury association nature.",
        "It does not stack retrograde and dignity multipliers.",
        "It does not apply special corner/junction rules, Latta, or classical natal-severity translation.",
        "It does not resolve absolute cardinal orientation.",
        "It is not yet prospectively validated for financial use.",
    ]
    for item in blocked:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Bottom line",
        "Use the Chakra Lab to organize and compare source-profiled evidence. Keep all market interpretation in a separately validated research workflow.",
        fill=RED_LIGHT,
        border=RED,
        label_color=RED,
    )
    doc.add_heading("Source and implementation references", level=2)
    references = [
        "Gann Astro Desk 0.10.9: src/views/ChakraLabWorkspace.tsx.",
        "Snapshot contract: sbc/chakra_lab.py, SBC_CHAKRA_LAB_SNAPSHOT_V1.",
        "Vedha engine and normalized ledger: sbc/vedha.py.",
        "Executable profile: configs/sbc/vedha/phaladeepika_editor_vedha_guidance_v1.yaml.",
        "81-cell profile: configs/sbc/grids/sbc_81_rotation_normalized_partial_v1.yaml.",
        "Program audit: docs/sbc/VEDHA_GUIDANCE_AUDIT.md.",
        "Primary rule witness: Phaladeepika 1937 editor supplement, PDF pages 347-351, printed pages 310-314.",
        "Secondary geometry/motion cross-check: Sanjay Rath, Crux of Vedic Astrology, PDF pages 21-22, printed pages 10-11.",
    ]
    for reference in references:
        add_bullet(doc, reference)
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.paragraph_format.space_before = Pt(20)
    final_run = final.add_run("END OF MANUAL  |  VERSION 1.0")
    set_run_font(final_run, 10, MUTED, bold=True)

    MANUAL_ROOT.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


def main() -> None:
    MANUAL_ROOT.mkdir(parents=True, exist_ok=True)
    ASSET_ROOT.mkdir(parents=True, exist_ok=True)
    if QA_ROOT.exists():
        shutil.rmtree(QA_ROOT)
    create_quick_start_infographic()
    create_annotated_screenshot()
    create_ray_geometry_infographic()
    build_document()
    print(OUTPUT_DOCX)
    print(QUICK_START_PNG)
    print(ANNOTATED_SCREEN_PNG)
    print(RAY_GEOMETRY_PNG)


if __name__ == "__main__":
    main()
